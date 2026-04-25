"""
extract_cotizacion.py
Extrae datos clave de cotizaciones Melectra (.docx).

Estructura esperada del documento (líneas no vacías):
  1  → COTIZACIÓN No. XXX
  2  → Fecha (con ciudad)
  3  → Saludo / título (Señor, Ingeniero, etc.)
  4  → Nombre de la persona
  5+ → Empresa (posiblemente precedida por cargo)
  …  → Email(s), Teléfono(s), ASUNTO

Campos extraídos:
  - numero        → cuerpo del doc (sin espacios) y/o nombre de archivo
  - nombre        → 4ª línea no vacía del documento
  - empresa       → 5ª+ línea no vacía, saltando cargo/ciudad/contacto
  - telefono      → todos los teléfonos (Móvil, Cel, Teléfono)
  - correo        → todos los correos encontrados
  - servicio      → después de "ASUNTO:" (párrafos continuos hasta línea vacía)
  - valor_total   → fila resumen o valor máximo en la columna de valor
  - fecha         → fecha al inicio del documento, convertida a DD/MM/YYYY
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import Table

from models import DatosCotizacion


# ---------------------------------------------------------------------------
# Utilidades de fecha
# ---------------------------------------------------------------------------

MESES_ES = {
    "enero": 1, "febrero": 2, "marzo": 3, "abril": 4,
    "mayo": 5, "junio": 6, "julio": 7, "agosto": 8,
    "septiembre": 9, "setiembre": 9, "octubre": 10,
    "noviembre": 11, "diciembre": 12,
}

RE_FECHA_NUM = re.compile(r"(\d{1,2})\s+de\s+(\w+)\s+del?\s+(\d{4})", re.IGNORECASE)


def fecha_a_ddmmyyyy(texto: str | None) -> str | None:
    """Convierte '4 de abril del 2026' a '04/04/2026'."""
    if not texto:
        return None
    m = RE_FECHA_NUM.search(texto)
    if not m:
        return texto
    dia = int(m.group(1))
    mes_nombre = m.group(2).lower()
    anio = int(m.group(3))
    mes = MESES_ES.get(mes_nombre)
    if mes is None:
        return texto
    return f"{dia:02d}/{mes:02d}/{anio}"


# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

RE_NUMERO    = re.compile(r"COTIZACI[OÓ]N\s+No\.?\s*(.+)", re.IGNORECASE)
RE_TEL       = re.compile(
    r"(?:M[oó]vil|Movil|Cel(?:ular)?|Tel[eé]fono(?:\s+fijo)?)\s*:\s*(.+)",
    re.IGNORECASE,
)
RE_EMAIL     = re.compile(r"E-?mail\s*:\s*(.+)", re.IGNORECASE)
RE_ASUNTO    = re.compile(r"ASUNTO\s*:\s*(.+)", re.IGNORECASE)
RE_FECHA     = re.compile(r"\d{1,2}\s+de\s+\w+\s+del?\s+\d{4}", re.IGNORECASE)
RE_VALOR_HDR = re.compile(r"VALOR\s+TOTAL(?:\s+DEL\s+MES)?[,]?\s+ANTES\s+(?:DE(?:L)?\s+)?IVA", re.IGNORECASE)

# Extrae el número de cotización del nombre del archivo: "COT 040401-26SV-W ..."
RE_NUM_FILE  = re.compile(r"COT\s+([\w-]+)", re.IGNORECASE)

# Extrae la empresa del nombre del archivo: palabras entre el número y la
# primera palabra clave que inicia la descripción del servicio
RE_EMP_FILE  = re.compile(
    r"COT\s+[\w-]+\s+(.+?)"
    r"(?=\s+(?:PROYECTO|PROPUESTA|PRUEBA[S]?|PARA\b|CORREC|DIAGNÓS|"
    r"MANTENIMIENTO|DISEÑO|INSTALA|VALIDAC|VLF\b)|\s+-\s|\.\w+$|$)",
    re.IGNORECASE,
)

# Patrones que indican que una línea NO es la empresa
RE_CONTACTO  = re.compile(r"M[oó]vil|Movil|Cel|Tel[eé]fono|E-?mail|ASUNTO", re.IGNORECASE)
RE_CARGO     = re.compile(
    r"^(?:Jef[ae]\s|Coordinador\s|Profesional\s|Analista\s|Asistente\s"
    r"|Gerente\s|Director\s|Ingenier[oía]|Compras$)",
    re.IGNORECASE,
)
RE_CIUDAD    = re.compile(
    r"^(?:Medell[ií]n|Copacabana|Itag[uü][ií]|Bogot[áa]|Cali"
    r"|Barranquilla|Bucaramanga|Cartagena|Cúcuta|Pereira|Manizales)$",
    re.IGNORECASE,
)

# Filtros para excluir datos de contacto propios de Melectra
_MELECTRA_DOMAINS = {"melectra.com", "melectra.com.co"}
_MELECTRA_PHONES  = {"3232774518", "3052655310", "3113753455"}


# ---------------------------------------------------------------------------
# Helpers de valor
# ---------------------------------------------------------------------------

def _parse_raw_valor(raw: str) -> Optional[float]:
    """'$ 9.800.000' | '$9,800,000' | '1.700.000' → float"""
    limpio = re.sub(r"[\$\s]", "", raw)   # quita $ y espacios
    # Si hay coma Y punto: coma es miles → quitarla, punto es decimal
    # Si sólo hay punto(s): todos son miles (formato colombiano)
    if "," in limpio and "." in limpio:
        limpio = limpio.replace(",", "")
    elif "," in limpio:
        # podría ser decimal europeo o miles; si hay más de 3 dígitos tras coma → miles
        partes = limpio.split(",")
        limpio = limpio.replace(",", "") if len(partes[-1]) == 3 else limpio.replace(",", ".")
    else:
        # sólo puntos → todos son separadores de miles
        limpio = limpio.replace(".", "")
    try:
        return float(limpio)
    except ValueError:
        return None


def _format_valor(n: float) -> str:
    """9800000 → '$9,800,000'"""
    return f"${n:,.0f}"


# ---------------------------------------------------------------------------
# Extractor modular
# ---------------------------------------------------------------------------

class CotizacionExtractor:
    """
    Extrae campos de un documento .docx de cotización Melectra.

    Cada método `_extract_*` es reemplazable por una implementación LLM.
    """

    def extract(self, path: str | Path) -> DatosCotizacion:
        path = Path(path)
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs]
        tables = doc.tables
        non_empty = [t for t in paragraphs if t]

        datos = DatosCotizacion()
        datos.numero      = self._extract_numero(paragraphs, path)
        datos.nombre      = self._extract_nombre(non_empty)
        datos.empresa     = self._extract_empresa(non_empty, path)

        # Si no se encontró empresa por posición, intentar extraerla del
        # mismo párrafo del nombre (ej: "Francisco Torres \nAMTEX")
        if not datos.empresa and datos.nombre:
            filtradas = [
                t for t in non_empty if not re.match(r"^[\d.\s]+$", t)
            ]
            if len(filtradas) >= 4:
                nombre_line = filtradas[3]
                partes = nombre_line.split("\n")
                if len(partes) > 1:
                    datos.empresa = partes[1].strip() or None

        datos.telefono    = self._extract_telefono(paragraphs)
        datos.correo      = self._extract_correo(paragraphs)
        datos.servicio    = self._extract_servicio(paragraphs)
        datos.valor_total = self._extract_valor_total(tables)
        datos.fecha       = self._extract_fecha(paragraphs)
        return datos

    # ------------------------------------------------------------------
    # Helpers de filename
    # ------------------------------------------------------------------

    def _numero_from_filename(self, path: Path) -> Optional[str]:
        m = RE_NUM_FILE.search(path.stem)
        if m:
            return re.sub(r"\s+", "", m.group(1))
        return None

    def _empresa_from_filename(self, path: Path) -> Optional[str]:
        m = RE_EMP_FILE.search(path.stem)
        if m:
            return m.group(1).strip()
        return None

    # ------------------------------------------------------------------
    # Campos individuales
    # ------------------------------------------------------------------

    def _extract_numero(self, paragraphs: list[str], path: Optional[Path] = None) -> Optional[str]:
        """
        Busca en el cuerpo del documento y en el nombre del archivo.
        Elimina espacios del resultado para que quede todo junto.
        """
        body = None
        for text in paragraphs:
            m = RE_NUMERO.search(text)
            if m:
                body = re.sub(r"\s+", "", m.group(1).strip())
                break

        filename = self._numero_from_filename(path) if path else None
        return body or filename

    def _extract_nombre(self, non_empty: list[str]) -> Optional[str]:
        """4ª línea no vacía del documento (índice 3).

        Filtra líneas que sean solo números/puntos (ej. "26.") y toma
        solo la primera parte si hay salto de línea embebido.
        """
        filtradas = [
            t for t in non_empty
            if not re.match(r"^[\d.\s]+$", t)
        ]
        if len(filtradas) >= 4:
            nombre = filtradas[3].strip()
            # Si hay salto de línea embebido, tomar solo la primera parte
            nombre = nombre.split("\n")[0].strip()
            # Limpiar puntuación final
            nombre = nombre.rstrip(". ")
            return nombre or None
        return None

    def _extract_empresa(self, non_empty: list[str], path: Optional[Path] = None) -> Optional[str]:
        """
        Busca la empresa a partir de la 5ª línea no vacía (índice 4).
        Salta líneas que sean cargo, ciudad, contacto o solo números.
        Compara con el nombre del archivo como referencia.
        """
        empresa_file = self._empresa_from_filename(path) if path else None

        filtradas = [
            t for t in non_empty
            if not re.match(r"^[\d.\s]+$", t)
        ]

        for j in range(4, min(len(filtradas), 10)):
            linea = filtradas[j].strip()
            if not linea:
                continue
            if RE_CONTACTO.search(linea):
                break
            if RE_ASUNTO.search(linea):
                break
            if RE_CARGO.match(linea):
                continue
            if RE_CIUDAD.match(linea):
                continue
            # Si hay referencia del filename, preferir la que coincida
            if empresa_file and empresa_file.upper() in linea.upper():
                return linea
            # Primera línea que no sea cargo/ciudad/contacto → empresa
            return linea

        # Fallback: usar el nombre del archivo
        return empresa_file

    @staticmethod
    def _is_melectra_phone(valor: str) -> bool:
        """Detecta teléfonos propios de Melectra por número conocido."""
        normalizado = re.sub(r"[\s\-]", "", valor)
        return normalizado in _MELECTRA_PHONES

    def _extract_telefono(self, paragraphs: list[str]) -> Optional[str]:
        """Extrae TODOS los teléfonos encontrados, excluyendo los de Melectra."""
        encontrados: list[str] = []
        for text in paragraphs:
            for m in RE_TEL.finditer(text):
                valor = m.group(1).strip()
                if valor and not self._is_melectra_phone(valor):
                    encontrados.append(valor)
        return ", ".join(encontrados) if encontrados else None

    def _extract_correo(self, paragraphs: list[str]) -> Optional[str]:
        """Extrae TODOS los correos encontrados, excluyendo los de Melectra."""
        encontrados: list[str] = []
        for text in paragraphs:
            m = RE_EMAIL.search(text)
            if m:
                raw = m.group(1).strip()
                partes = re.split(r"[;,\s]+(?:y\s+)?(?=\S+@)", raw)
                for parte in partes:
                    parte = parte.strip().strip(";,–- .")
                    if parte and "@" in parte:
                        dominio = parte.split("@")[-1].lower()
                        if dominio not in _MELECTRA_DOMAINS:
                            encontrados.append(parte)
        return ", ".join(encontrados) if encontrados else None

    def _extract_servicio(self, paragraphs: list[str]) -> Optional[str]:
        """
        Captura todo el texto del párrafo ASUNTO y sus continuaciones
        hasta encontrar una línea vacía.
        """
        for i, text in enumerate(paragraphs):
            m = RE_ASUNTO.search(text)
            if m:
                partes = [m.group(1).strip()]
                for siguiente in paragraphs[i + 1:]:
                    if not siguiente:
                        break
                    partes.append(siguiente)
                return " ".join(p for p in partes if p)
        return None

    def _extract_fecha(self, paragraphs: list[str]) -> Optional[str]:
        """Extrae la fecha y la convierte a DD/MM/YYYY."""
        for text in paragraphs:
            m = RE_FECHA.search(text)
            if m:
                return fecha_a_ddmmyyyy(m.group(0).strip())
        return None

    def _extract_valor_total(self, tables: list[Table]) -> Optional[str]:
        """
        Busca en la última tabla con columna de valor:
        1. Si hay fila resumen "VALOR TOTAL ANTES DE IVA" → usar ese valor.
        2. Si no → tomar el valor más alto de la columna.
        Limpia '=' al final del valor.
        """
        max_val: Optional[float] = None
        resumen_val: Optional[float] = None

        for table in tables:
            if not table.rows:
                continue
            header_cells = [c.text.strip() for c in table.rows[0].cells]
            col_idx = next(
                (i for i, h in enumerate(header_cells) if RE_VALOR_HDR.search(h)),
                None,
            )
            if col_idx is None:
                continue

            for row in table.rows[1:]:
                if col_idx >= len(row.cells):
                    continue
                raw = row.cells[col_idx].text.strip().rstrip("=").strip()
                if not raw:
                    continue
                n = _parse_raw_valor(raw)
                if n is None:
                    continue

                # Detectar fila resumen (primera celda contiene "VALOR TOTAL")
                primera_celda = row.cells[0].text.strip().upper()
                if "VALOR TOTAL" in primera_celda:
                    resumen_val = n
                elif max_val is None or n > max_val:
                    max_val = n

        resultado = resumen_val if resumen_val is not None else max_val
        return _format_valor(resultado) if resultado is not None else None
